import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_heading(doc, text, level):
    heading = doc.add_heading(text, level=level)
    return heading

def add_paragraph(doc, text):
    p = doc.add_paragraph(text)
    return p

def main():
    doc = Document()
    
    # Title
    title = doc.add_heading('Travel Booking Platform Project Documentation', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Introduction
    add_heading(doc, '1. Project Overview', level=1)
    add_paragraph(doc, 'YatraBook is a full-stack, production-level travel booking platform for India. It allows users to search and book Trains, Flights, and Buses with a premium, futuristic dark-themed UI. The platform is built using the MERN stack (MongoDB, Express, React, Node.js) along with Tailwind CSS and Framer Motion.')
    
    add_paragraph(doc, 'Key Features:')
    doc.add_paragraph('Multi-modal Search (Trains, Flights, Buses)', style='List Bullet')
    doc.add_paragraph('Smart Filters (Price, duration, departure time, class)', style='List Bullet')
    doc.add_paragraph('Booking Simulation with seat selection, payment flow, and PNR generation', style='List Bullet')
    doc.add_paragraph('Waitlist System and Recommendations for cheapest/fastest routes', style='List Bullet')
    
    # Project Structure
    add_heading(doc, '2. Project Structure', level=1)
    add_paragraph(doc, 'The project is divided into two main parts: client (React frontend) and server (Express backend).')
    
    # Client Structure
    add_heading(doc, '3. Frontend (client/)', level=1)
    add_paragraph(doc, 'The frontend is built with React 18, Vite, Tailwind CSS, and Framer Motion. It resides in the `client/` directory.')
    
    add_heading(doc, '3.1 Pages (client/src/pages/)', level=2)
    doc.add_paragraph('HomePage.jsx: The main landing page with search forms and general platform information.', style='List Bullet')
    doc.add_paragraph('SearchPage.jsx: Displays search results for flights, trains, and buses, complete with filtering options.', style='List Bullet')
    doc.add_paragraph('BookingPage.jsx: Handles the booking process, including seat selection and payment flow.', style='List Bullet')
    doc.add_paragraph('DashboardPage.jsx: User profile page showing booking history and recent searches.', style='List Bullet')
    doc.add_paragraph('LoginPage.jsx & SignupPage.jsx: User authentication pages.', style='List Bullet')
    
    add_heading(doc, '3.2 Components (client/src/components/)', level=2)
    doc.add_paragraph('layout/: Reusable layout components like Navbar, Footer, etc.', style='List Bullet')
    doc.add_paragraph('ui/: Reusable UI components like buttons, inputs, modals, etc.', style='List Bullet')
    doc.add_paragraph('PrintTicket.jsx: A component to render and print booked tickets.', style='List Bullet')
    
    add_heading(doc, '3.3 State & Logic (client/src/)', level=2)
    doc.add_paragraph('features/: Redux or Context state slices for managing global application state.', style='List Bullet')
    doc.add_paragraph('lib/: Utility functions and API clients.', style='List Bullet')
    doc.add_paragraph('styles/: Global CSS stylesheets.', style='List Bullet')
    doc.add_paragraph('App.jsx & main.jsx: Application entry points.', style='List Bullet')

    # Server Structure
    add_heading(doc, '4. Backend (server/)', level=1)
    add_paragraph(doc, 'The backend is built with Node.js, Express, and MongoDB (Mongoose). It resides in the `server/` directory.')
    
    add_heading(doc, '4.1 Models (server/src/models/)', level=2)
    add_paragraph(doc, 'Mongoose schemas defining the structure of MongoDB collections.')
    doc.add_paragraph('User.js: User profile, authentication details, and booking history references.', style='List Bullet')
    doc.add_paragraph('Flight.js, Train.js, Bus.js: Models defining schedules, routes, seat availability, and pricing for each transport mode.', style='List Bullet')
    doc.add_paragraph('Booking.js: Stores booking details like PNR, user ID, selected seats, total price, and payment status.', style='List Bullet')
    
    add_heading(doc, '4.2 Controllers (server/src/controllers/)', level=2)
    add_paragraph(doc, 'Handles the business logic for incoming API requests.')
    doc.add_paragraph('auth.controller.js: Registration, login, and JWT generation.', style='List Bullet')
    doc.add_paragraph('booking.controller.js: Creating new bookings, fetching user bookings, and processing payments.', style='List Bullet')
    doc.add_paragraph('flight.controller.js, train.controller.js, bus.controller.js: Fetching and filtering transport data.', style='List Bullet')
    doc.add_paragraph('recommend.controller.js: Logic for finding cheapest and fastest recommendations.', style='List Bullet')
    doc.add_paragraph('user.controller.js: Fetching and updating user profiles.', style='List Bullet')
    
    add_heading(doc, '4.3 Routes (server/src/routes/)', level=2)
    add_paragraph(doc, 'Express routers that map endpoints to controller functions.')
    
    add_heading(doc, '4.4 Middleware (server/src/middleware/)', level=2)
    add_paragraph(doc, 'Functions that intercept requests before they reach controllers, such as JWT authentication and error handling.')

    add_heading(doc, '4.5 Config & Utilities (server/src/)', level=2)
    doc.add_paragraph('config/: Database connection setup and environment variable loading.', style='List Bullet')
    doc.add_paragraph('utils/: Helper functions (e.g., error formatters, date manipulators).', style='List Bullet')
    doc.add_paragraph('validators/: Request payload validation (e.g., using Joi or express-validator).', style='List Bullet')
    doc.add_paragraph('server.js & app.js: Entry points configuring Express application and starting the HTTP server.', style='List Bullet')

    # Conclusion
    add_heading(doc, '5. Summary of Functioning', level=1)
    add_paragraph(doc, '1. A user visits the React frontend and enters a search query on the HomePage.')
    add_paragraph(doc, '2. The React app sends a GET request to the relevant Express backend endpoint (e.g., /api/flights).')
    add_paragraph(doc, '3. The Express controller processes the request, filters data from the MongoDB database using Mongoose models, and returns JSON.')
    add_paragraph(doc, '4. The React frontend updates its state and displays results on the SearchPage.')
    add_paragraph(doc, '5. The user selects a route, logs in (handled via auth controller & JWT), and proceeds to BookingPage.')
    add_paragraph(doc, '6. Upon payment simulation, a new Booking document is created in MongoDB, and the user can view/print their ticket.')

    doc.save('Travel_Booking_Platform_Documentation.docx')
    print("Documentation generated successfully.")

if __name__ == '__main__':
    main()
